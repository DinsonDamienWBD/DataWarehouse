import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_resume():
    doc = Document()
    
    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Header ---
    name = doc.add_paragraph()
    name_run = name.add_run('Athulya T. Manayil')
    name_run.bold = True
    name_run.font.size = Pt(20)
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    title = doc.add_paragraph()
    title_run = title.add_run('Sourcing & Procurement Engineer')
    title_run.font.size = Pt(14)
    title_run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    contact = doc.add_paragraph('Tokyo, Japan (Suginami-ku) | +81 80-1033-5414 | athulyatmanayil@gmail.com\nhttps://www.linkedin.com/in/athulya-manayil-721402218')
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph() # Spacer

    # --- Visa Status ---
    doc.add_heading('VISA STATUS', level=2)
    p_visa = doc.add_paragraph()
    p_visa.add_run('Current Status: ').bold = True
    p_visa.add_run('Dependent Visa (Authorized for part-time work).\n')
    p_visa.add_run('Availability: ').bold = True
    p_visa.add_run('Available for immediate start. ')
    p_visa.add_run('Requires Visa sponsorship/documentation support').bold = True
    p_visa.add_run(' for full-time permanent employment.')

    # --- Summary ---
    doc.add_heading('PROFESSIONAL SUMMARY', level=2)
    doc.add_paragraph(
        "Results-oriented Sourcing Engineer with over 5 years of experience in the electronics and ODM "
        "manufacturing sectors. Proven track record in strategic procurement, global supplier relationship management, "
        "and cost reduction. Currently based in Tokyo with Intermediate+ Japanese proficiency and successfully adapted "
        "to the Japanese work environment. Expert in contract negotiation, risk mitigation, and inventory control using "
        "ERP systems (Microsoft Nav/Dynamics)."
    )

    # --- Core Competencies ---
    doc.add_heading('CORE COMPETENCIES', level=2)
    competencies = [
        ("Strategic Sourcing", "Supplier Evaluation, Cost Analysis, Contract Negotiation, Global Procurement."),
        ("Supply Chain Management", "Inventory Control, Demand Forecasting, Risk Mitigation, Lead Time Reduction."),
        ("Technical Skills", "ERP Systems (Microsoft Nav/Dynamics), Microsoft Office Suite, Electronics Component Knowledge (Active/Passive)."),
        ("Languages", "English (Fluent), Hindi (Fluent), Japanese (Intermediate+).")
    ]
    for category, details in competencies:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{category}: ").bold = True
        p.add_run(details)

    # --- Experience ---
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)

    # Meiko Kids
    p = doc.add_paragraph()
    p.add_run('MEIKO KIDS').bold = True
    p.add_run(' | Tokyo, Japan')
    p_sub = doc.add_paragraph()
    p_sub.add_run('English Coach / Staff').italic = True
    p_sub.add_run(' | Aug 2025 – Present').italic = True
    doc.add_paragraph('Current role demonstrating successful integration into the Japanese workforce.')
    
    bullets_job1 = [
        ("Japanese Work Culture", "Successfully adapted to Japanese professional standards, including workplace etiquette, punctuality, and team coordination."),
        ("Communication", "Collaborating daily with Japanese staff and management, utilizing intermediate Japanese skills to ensure smooth operations and support student activities."),
        ("Instruction", "Providing English language training and activity support for elementary school children, fostering a bilingual learning environment.")
    ]
    for t, d in bullets_job1:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{t}: ").bold = True
        p.add_run(d)

    # SFO
    p = doc.add_paragraph()
    p.add_run('\nSFO TECHNOLOGIES (NeST Group)').bold = True
    p.add_run(' | Kerala, India')
    p_sub = doc.add_paragraph()
    p_sub.add_run('Sourcing Engineer (SCM Department)').italic = True
    p_sub.add_run(' | Dec 2018 – Nov 2022').italic = True
    doc.add_paragraph('Flagship company of NeST Group providing ODM Plus solutions to Aerospace, Defense, and Communications industries ($300M+ Revenue).')
    
    bullets_job2 = [
        ("Strategic Procurement", "Managed the end-to-end procurement of Active and Passive electronic components. Successfully secured favorable pricing and terms through strategic negotiations with global suppliers."),
        ("Cost Optimization", "Consistently achieved substantial cost savings by consolidating suppliers and optimizing purchasing agreements without compromising quality."),
        ("Contract Management", "Navigated complex contracts, enforcing strict adherence to specifications. Managed Minimum Order Quantities (MOQ) and Non-Cancellable Non-Returnable (NCNR) validations to prevent overstocking."),
        ("Risk Mitigation", "Proactively identified supply chain risks, such as lead time fluctuations and component shortages, and implemented mitigation strategies to ensure production continuity."),
        ("Operational Leadership", "Mentored commodity buyers in resolving day-to-day operational challenges. Maintained high data integrity within the ERP system to ensure accurate demand planning and forecasting."),
        ("Supplier Relations", "Conducted regular performance reviews and fostered strong partnerships with key vendors to ensure priority support and material availability.")
    ]
    for t, d in bullets_job2:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{t}: ").bold = True
        p.add_run(d)

    # Elmactics
    p = doc.add_paragraph()
    p.add_run('\nELMACTICS ENTERPRISES').bold = True
    p.add_run(' | Kerala, India')
    p_sub = doc.add_paragraph()
    p_sub.add_run('Technical Assistant (Procurement & Sales)').italic = True
    p_sub.add_run(' | Sep 2017 – Dec 2018').italic = True
    doc.add_paragraph('Sourcing and distribution company for engineering laboratory tools and devices.')
    
    bullets_job3 = [
        ("Tender Management", "Managed the comprehensive lifecycle of customer quotations and tenders, ensuring 100% accuracy and timely submission aligned with client requirements."),
        ("Vendor Negotiation", "Negotiated with suppliers to procure engineering tools at optimal price points, directly contributing to improved project margins."),
        ("Process Improvement", "Maintained meticulous records and database systems, streamlining the documentation process for sales and procurement.")
    ]
    for t, d in bullets_job3:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{t}: ").bold = True
        p.add_run(d)

    # --- Education ---
    doc.add_heading('EDUCATION', level=2)
    p = doc.add_paragraph()
    p.add_run('Bachelor of Technology (B.Tech) in Electronics and Communication').bold = True
    doc.add_paragraph('KMEA Engineering College, Kerala, India | 2013 – 2017')

    # --- Skills ---
    doc.add_heading('TECHNICAL SKILLS', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Software: ').bold = True
    p.add_run('Microsoft Dynamics (NAV), ERP Systems, Microsoft Office (Excel, Word, PowerPoint).')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Operating Systems: ').bold = True
    p.add_run('Windows.')

    # --- SAVE TO DOWNLOADS LOGIC ---
    # Automatically detect the user's home directory and find "Downloads"
    downloads_path = Path.home() / "Downloads"
    filename = "Athulya_Manayil_Resume.docx"
    full_save_path = downloads_path / filename

    try:
        doc.save(full_save_path)
        print(f"✅ Success! File saved to: {full_save_path}")
    except Exception as e:
        print(f"❌ Error saving to Downloads folder: {e}")
        print("Attempting to save to current folder instead...")
        doc.save(filename)
        print(f"Saved to current folder: {filename}")

if __name__ == "__main__":
    create_resume()